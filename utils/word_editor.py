import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
import datetime
import random
from utils.const_data import *

def create_report(number, info, owner):
    doc = docx.Document()
    dt = datetime.datetime.now()

    doc.add_picture('logo.png', width = docx.shared.Cm(3))
    t = doc.add_paragraph('Федеральное агентство по рыболовству')
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')

    t = doc.add_paragraph(f'Отчет №{number} от {str(dt).split(".")[0]}')
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT

    items = info

    table = doc.add_table(1, len(items[0]))
    table.style = 'Light Shading Accent 1'
    head_cells = table.rows[0].cells

    for i, item in enumerate(['судно', 'держатель', 'платформа', 'рыба', 'регион']):
        p = head_cells[i].paragraphs[0]
        p.add_run(item).bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for row in items:
        cells = table.add_row().cells
        for i, item in enumerate(row):
            cells[i].text = str(item)
            if i == 2:
                cells[i].paragraphs[0].runs[0].font.name = 'Arial'
    doc.add_paragraph('')
    t = doc.add_paragraph(f'{owner} ______________')
    t.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.save(DATA_PATH + 'picture.docx')
    return DATA_PATH + 'picture.docx'

def create_data_list(data_df: pd.DataFrame):
    data_list = list(())
    data_row = list()
    for index, row in data_df.iterrows():
        data_row = [row['id_ves'], row['id_own'], row['Name_Plat'], row['fish'], row['Region_Plat']]
        data_list.append(data_row)
    return data_list

